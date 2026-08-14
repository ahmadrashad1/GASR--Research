import sys
from docx import Document
from docx.shared import Pt

md_path = sys.argv[1]
docx_path = sys.argv[2]

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

in_code = False
code_lines = []

with open(md_path, 'r', encoding='utf-8') as f:
    for raw in f:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # end code block
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue
        stripped = line.strip()
        if stripped == '':
            doc.add_paragraph()
            continue
        # headings
        if stripped.startswith('#'):
            # count leading #
            i = 0
            while i < len(stripped) and stripped[i] == '#':
                i += 1
            level = min(i, 4)
            text = stripped[i:].strip()
            if text == '':
                continue
            doc.add_heading(text, level=level)
            continue
        # unordered list
        if stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
            continue
        # numbered list
        if stripped[:2].isdigit() and stripped[2:3] == '.':
            doc.add_paragraph(stripped[4:], style='List Number')
            continue
        # regular paragraph
        doc.add_paragraph(line)

try:
    doc.save(docx_path)
    print('DOCX saved:', docx_path)
except Exception as e:
    print('Failed to save DOCX:', e)
    sys.exit(1)
